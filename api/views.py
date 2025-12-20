# backend/api/views.py
import io
import re
import spacy
from spacy.pipeline import EntityRuler
from sentence_transformers import SentenceTransformer, util
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, viewsets
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.parsers import MultiPartParser, FormParser
import PyPDF2
from docx import Document
from django.conf import settings
import google.generativeai as genai
from .models import Resume, Analysis
from .serializers import UserSerializer, ResumeSerializer
from django.contrib.auth.models import User
import language_tool_python
import json
import os

# --- Load Models and Skills ---
nlp = None
similarity_model = None
lang_tool = None
SKILL_KEYWORDS = []

# Fallback skills in case AI fails or key is missing
FALLBACK_ROLE_SKILLS = {
    "Software Engineer": ["Python", "Java", "C++", "SQL", "Git", "Data Structures", "Algorithms", "System Design"],
    "Frontend Developer": ["HTML", "CSS", "JavaScript", "React", "Angular", "Vue", "TypeScript", "Redux", "Responsive Design"],
    "Backend Developer": ["Node.js", "Python", "Java", "Django", "Flask", "Spring Boot", "SQL", "NoSQL", "API Design"],
    "Full Stack Developer": ["HTML", "CSS", "JavaScript", "React", "Node.js", "Python", "SQL", "MongoDB", "Git", "AWS"],
    "Data Scientist": ["Python", "R", "SQL", "Machine Learning", "Pandas", "NumPy", "Scikit-learn", "TensorFlow", "Data Visualization"],
    "Machine Learning Engineer": ["Python", "TensorFlow", "PyTorch", "Deep Learning", "NLP", "Computer Vision", "MLOps", "SQL"],
    "DevOps Engineer": ["Linux", "AWS", "Azure", "Docker", "Kubernetes", "Jenkins", "Terraform", "CI/CD", "Bash Scripting"],
    "Project Manager": ["Agile", "Scrum", "JIRA", "Communication", "Risk Management", "Leadership", "Planning", "Stakeholder Management"],
    "UI/UX Designer": ["Figma", "Adobe XD", "Sketch", "Prototyping", "User Research", "Wireframing", "Usability Testing", "Visual Design"],
    "QA Engineer": ["Selenium", "Java", "Python", "Test Automation", "Manual Testing", "JIRA", "SQL", "API Testing"],
    "Business Analyst": ["SQL", "Excel", "Tableau", "Power BI", "Data Analysis", "Requirements Gathering", "Communication", "Documentation"]
}

try:
    skills_file_path = os.path.join(os.path.dirname(__file__), 'skills.json')
    if os.path.exists(skills_file_path):
        with open(skills_file_path, 'r', encoding='utf-8') as f:
            SKILL_KEYWORDS = json.load(f)
    
    nlp = spacy.load("en_core_web_lg")

    if SKILL_KEYWORDS:
        patterns = [{"label": "SKILL", "pattern": skill} for skill in SKILL_KEYWORDS]
        if "entity_ruler" not in nlp.pipe_names:
            ruler = nlp.add_pipe("entity_ruler", before="ner", config={"phrase_matcher_attr": "LOWER"})
            ruler.add_patterns(patterns)
    
    similarity_model = SentenceTransformer('all-MiniLM-L6-v2')
    lang_tool = language_tool_python.LanguageTool('en-US')
    print("AI Models loaded successfully.")

except Exception as e:
    print(f"Error loading AI models: {e}")
    nlp = None

# --- Helpers ---
def extract_text_from_pdf(file_obj):
    text = ""
    try:
        pdf_reader = PyPDF2.PdfReader(file_obj)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text: text += page_text + "\n"
    except Exception as e:
        print(f"PDF Error: {e}")
    return text if text else None

def extract_text_from_docx(file_obj):
    try:
        doc = Document(file_obj)
        full_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text: full_text.append(" | ".join(row_text))
        for para in doc.paragraphs:
            if para.text.strip(): full_text.append(para.text.strip())
        return "\n".join(full_text)
    except Exception as e:
        print(f"DOCX Error: {e}")
        return None

# --- Views ---
class HelloApiView(APIView):
    permission_classes = [AllowAny]
    def get(self, request): return Response({'message': 'Hello!'})

class RegisterView(APIView):
    permission_classes = [AllowAny]
    def post(self, request):
        serializer = UserSerializer(data=request.data)
        if serializer.is_valid():
            User.objects.create_user(username=serializer.validated_data['username'], email=serializer.validated_data.get('email', ''), password=serializer.validated_data['password'])
            return Response({'username': serializer.validated_data['username']}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class EnhanceWithAIView(APIView):
    permission_classes = [IsAuthenticated]
    def post(self, request):
        text_to_enhance = request.data.get('text')
        prompt_override = request.data.get('prompt_override')
        if not text_to_enhance: return Response({'error': 'No text'}, status=400)
        try:
            if not settings.GEMINI_API_KEY: return Response({'error': 'No API Key'}, status=500)
            genai.configure(api_key=settings.GEMINI_API_KEY)
            model = genai.GenerativeModel('gemini-pro')
            prompt = prompt_override if prompt_override else f"Rewrite this resume bullet point to be professional: \"{text_to_enhance}\""
            response = model.generate_content(prompt)
            return Response({'enhanced_text': response.text}, status=200)
        except Exception as e:
            return Response({'error': str(e)}, status=500)

class ResumeViewSet(viewsets.ModelViewSet):
    serializer_class = ResumeSerializer
    permission_classes = [IsAuthenticated]
    def get_queryset(self): return Resume.objects.filter(user=self.request.user).order_by('-updated_at')
    def perform_create(self, serializer): serializer.save(user=self.request.user)

class ResumeAnalysisView(APIView):
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request, *args, **kwargs):
        if not nlp or not similarity_model:
             return Response({"success": False, "error": "AI models failed to load."}, status=503)

        resume_file = request.FILES.get('resume_file')
        job_role = request.data.get('job_role')
        job_description = request.data.get('job_description', '')

        if not resume_file or not job_role:
            return Response({"success": False, "error": "Missing file or job role."}, status=400)

        # 1. Extract Text
        try:
            file_stream = io.BytesIO(resume_file.read())
            if resume_file.name.lower().endswith('.pdf'):
                resume_text = extract_text_from_pdf(file_stream)
            elif resume_file.name.lower().endswith('.docx'):
                resume_text = extract_text_from_docx(file_stream)
            else: return Response({"success": False, "error": "Invalid file type"}, status=400)
            
            if not resume_text or not resume_text.strip():
                return Response({"success": False, "error": "Empty file"}, status=400)
        except Exception as e:
            return Response({"success": False, "error": f"File error: {e}"}, status=500)

       # --- Basic Info Extraction ---

        extracted_email = None

        extracted_phone = None

        extracted_name = None

        try:

            # 1. Find Name

            for line in resume_text.split('\n'):

                line = line.strip()

                if line and '@' not in line and not re.search(r'\d{5,}', line) and len(line.split()) < 5:

                    if not re.search(r'(experience|education|skills|projects|summary|profile)', line, re.IGNORECASE):

                        extracted_name = line

                        break

            

            # 2. Find Email

            email_pattern_clean = r'(?<=[^a-zA-Z0-9])([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b'

            emails_clean = re.findall(email_pattern_clean, resume_text, re.IGNORECASE)

            

            if emails_clean:

                extracted_email = emails_clean[0]

            else:

                email_pattern_merged = r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b'

                text_for_email = re.sub(r'[\s\|]+', '', resume_text)

                emails_merged = re.findall(email_pattern_merged, text_for_email, re.IGNORECASE)

                

                if emails_merged:

                    potential_email = emails_merged[0]

                    if extracted_name:

                        name_no_space = re.sub(r'\s+', '', extracted_name).lower()

                        if potential_email.lower().startswith(name_no_space):

                             email_part = potential_email[len(name_no_space):]

                             if re.fullmatch(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email_part, re.IGNORECASE):

                                 extracted_email = email_part

                             else:

                                 extracted_email = potential_email

                        else:

                             extracted_email = potential_email

                    else:

                        extracted_email = potential_email

            

            # 3. Find Phone

            text_no_space = re.sub(r'\s+', '', resume_text)

            phone_pattern_strict = r'(\+?\d{9,15})\b'

            phone_match = re.search(phone_pattern_strict, text_no_space)

            if phone_match:

                digits = phone_match.group(0).replace('+', '')[-10:]

                original_phone_pattern = r'[\D]*'.join(digits)

                original_match = re.search(original_phone_pattern, resume_text)

                if original_match:

                    extracted_phone = original_match.group(0).strip()

                else:

                    phone_pattern_orig = r'(\(?\+?\d{1,3}\)?[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'

                    phones_orig = re.findall(phone_pattern_orig, resume_text)

                    for phone in phones_orig:

                        phone_str = "".join(phone) if isinstance(phone, tuple) else phone

                        digits_only = re.sub(r'\D', '', phone_str)

                        if len(digits_only) >= 9:

                            extracted_phone = phone_str.strip()

                            break

                    if not extracted_phone:

                        extracted_phone = phone_match.group(0)

            

        except Exception as e:

            print(f"Regex error: {e}")

        # -----------------------------------

        # 3. Skill Extraction
        resume_skills_list = []
        try:
            doc = nlp(resume_text)
            unique_skills = set()
            for ent in doc.ents:
                if ent.label_ == "SKILL": unique_skills.add(ent.text.strip())
            if not unique_skills:
                for ent in doc.ents:
                    if ent.label_ in ["ORG", "PRODUCT", "LANGUAGE"]:
                        if len(ent.text.split()) < 4: unique_skills.add(ent.text.strip())
            resume_skills_list = sorted(list(set(s.capitalize() for s in unique_skills)))
        except Exception: pass

        # 4. Analysis & Scoring
        ats_score_role = 0
        ats_score_jd = None
        role_matching_skills = []
        role_missing_skills = []
        jd_matching_skills = []
        jd_missing_skills = []
        quality_feedback = []

        # A. General Quality Logic (Applied to Role Score)
        quality_score = 0
        sections = ['experience', 'education', 'skills', 'projects']
        found_sections = sum(1 for s in sections if re.search(r'\b'+s+r'\b', resume_text, re.IGNORECASE))
        quality_score += (found_sections * 10)
        
        if extracted_email and extracted_phone: quality_score += 10
        elif extracted_email or extracted_phone: quality_score += 5; quality_feedback.append("Contact info incomplete.")
        else: quality_feedback.append("Missing contact info.")

        if len(resume_skills_list) >= 5: quality_score += 10

        clean_grammar = re.sub(r'[^a-zA-Z0-9\s.,!?\'"-]', '', resume_text[:2000])
        errors = lang_tool.check(re.sub(r'\s+', ' ', clean_grammar))
        real_errors = [e for e in errors if not (e.ruleId == 'MORFOLOGIK_RULE_EN_US' and len(e.context.split()) < 3)]
        quality_score += max(0, 20 - len(real_errors))
        if real_errors: quality_feedback.append(f"Found {len(real_errors)} potential grammar issues.")
        quality_score = min(100, quality_score + 10)

        # B. Role-Based Analysis
        role_keywords = set()
        # 1. Try AI
        if settings.GEMINI_API_KEY:
            try:
                model = genai.GenerativeModel('gemini-pro')
                prompt = f"List top 20 technical skills and keywords for a '{job_role}' resume. Return ONLY comma-separated words. Do not include the job title itself."
                resp = model.generate_content(prompt)
                if resp.text:
                    items = re.split(r'[,\n\r•*-]+', resp.text)
                    for item in items:
                        clean = item.strip().lower().replace('skills', '').replace(job_role.lower(), '')
                        if len(clean) > 1: role_keywords.add(clean)
            except Exception: pass
        
        # 2. Use Fallback if AI failed or returned nothing
        if not role_keywords:
             fallback = FALLBACK_ROLE_SKILLS.get(job_role, ["communication", "teamwork", "problem solving"])
             for sk in fallback: role_keywords.add(sk.lower())

        resume_lower = set(s.lower() for s in resume_skills_list)
        for target in role_keywords:
            if target in resume_lower or any(target in s for s in resume_lower) or any(s in target for s in resume_lower):
                role_matching_skills.append(target.capitalize())
            else:
                role_missing_skills.append(target.capitalize())
        
        role_match_pct = (len(role_matching_skills) / len(role_keywords)) * 100 if role_keywords else 0
        ats_score_role = round((quality_score * 0.3) + (role_match_pct * 0.7)) # 70% Skills, 30% Quality

        # C. JD-Based Analysis
        if job_description.strip():
            jd_keywords = set()
            jd_doc = nlp(job_description)
            for ent in jd_doc.ents:
                if ent.label_ in ["SKILL", "ORG", "PRODUCT", "LANGUAGE"]: jd_keywords.add(ent.text.lower())
            
            for target in jd_keywords:
                if target in resume_lower or any(target in s for s in resume_lower) or any(s in target for s in resume_lower):
                    jd_matching_skills.append(target.capitalize())
                else:
                    jd_missing_skills.append(target.capitalize())

            jd_match_pct = (len(jd_matching_skills) / len(jd_keywords)) * 100 if jd_keywords else 0
            
            emb1 = similarity_model.encode(resume_text, convert_to_tensor=True)
            emb2 = similarity_model.encode(job_description, convert_to_tensor=True)
            semantic = util.cos_sim(emb1, emb2).item() * 100
            
            ats_score_jd = round((semantic * 0.5) + (jd_match_pct * 0.3) + (quality_score * 0.2))
            ats_score_jd = max(0, min(100, ats_score_jd))

        # Report
        summary_message = f"Analysis for {job_role}. Role Match: {ats_score_role}%."
        if ats_score_jd: summary_message += f" JD Match: {ats_score_jd}%."

        report = {
            "success": True,
            "job_role_selected": job_role,
            "name": extracted_name,
            "email": extracted_email,
            "phone": extracted_phone,
            "resume_skills": resume_skills_list,
            # Separated Lists
            "role_matching_skills": sorted(list(set(role_matching_skills))),
            "role_missing_skills": sorted(list(set(role_missing_skills))),
            "jd_matching_skills": sorted(list(set(jd_matching_skills))),
            "jd_missing_skills": sorted(list(set(jd_missing_skills))),
            # Scores
            "ats_score_general": ats_score_role, # Mapped for frontend compatibility
            "ats_score_role": ats_score_role,
            "ats_score_jd": ats_score_jd,
            "analysis_summary": summary_message,
            "quality_feedback": quality_feedback
        }
        
        try:
             Analysis.objects.create(user=request.user, job_role=job_role, resume_file=resume_file, ats_score_general=ats_score_role, ats_score_jd_match=ats_score_jd, analysis_result=report)
        except Exception as e:
            print(f"DB Save Error: {e}")

        return Response(report, status=200)